"""Sinh file Word bao cao chuong "Mo hinh hoc may" cho do an tot nghiep.

Hai muc:
  1. Quy trinh xay dung mo hinh hoc may (kem 2 bang dau vao / dau ra + 3 hinh)
  2. Danh gia mo hinh hoc may (kem 1 hinh)

Output: docs/chuong_mo_hinh_hoc_may.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
IMG_DIR = ROOT / "docs" / "anh"
OUT = ROOT / "docs" / "chuong_mo_hinh_hoc_may_v3.docx"

HEAD_COLOR = RGBColor(0x1A, 0x3A, 0x72)


def set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = HEAD_COLOR
        run.font.name = "Times New Roman"
    return h


def add_para(doc, text, *, bold=False, italic=False, align=None, size=12.5):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_image(doc, img_name, caption, width_cm=15.5):
    img_path = IMG_DIR / img_name
    if not img_path.exists():
        add_para(doc, f"[Thiếu ảnh: {img_name}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.font.name = "Times New Roman"
    cap_run.font.size = Pt(11.5)
    cap_run.italic = True


def add_table(doc, headers, rows, *, col_widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], "1A3A72")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    return table


# ============================================================================
# Khoi tao tai lieu
# ============================================================================
doc = Document()

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12.5)

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

# ============================================================================
# Muc 1 - Quy trinh xay dung mo hinh hoc may
# ============================================================================
add_heading(doc, "1. Quy trình xây dựng mô hình học máy", level=1)

add_para(
    doc,
    "Mô hình học máy của hệ thống có nhiệm vụ dự đoán giá trị cường độ tín hiệu nhận "
    "được (RSSI) trên đường truyền LoRa. Quy trình xây dựng mô hình được tổ chức "
    "thành bảy phần: thu thập dữ liệu khảo sát và xác định đặc trưng đầu vào, cách "
    "tính các đặc trưng địa hình và vật cản, phân tầng dữ liệu chống rò rỉ, kiến trúc "
    "học phần dư hai tầng, huấn luyện thuật toán Extra Trees, quy trình huấn luyện "
    "lại tự động qua hàng đợi Celery, và triển khai mô hình vào hệ thống phục vụ "
    "qua giao diện lập trình ứng dụng. Sơ đồ tổng thể được trình bày trong Hình 5.1.",
)

add_image(
    doc, "hinh_5_1.png", "Hình 5.1. Sơ đồ quy trình xây dựng mô hình học máy dự đoán RSSI LoRa"
)

# --- 1.1 Du lieu va dac trung ---
add_heading(doc, "1.1. Nguồn dữ liệu và đặc trưng đầu vào", level=2)

add_para(
    doc,
    "Nguồn dữ liệu huấn luyện được lấy từ bảng ts.survey_training trong cơ sở dữ liệu "
    "TimescaleDB của hệ thống. Mỗi bản ghi tương ứng với một gói tin mà thiết bị LoRa "
    "gửi thành công về một gateway, kèm theo toạ độ địa lý, dấu thời gian, hệ số trải "
    "phổ và giá trị RSSI đo được. Tổng cộng có 14 017 mẫu hợp lệ được sử dụng trong "
    "lần huấn luyện gần nhất.",
)
add_para(
    doc,
    "Mô hình sử dụng 21 đặc trưng đầu vào, được chia thành ba nhóm chính: nhóm đặc trưng "
    "vị trí và truyền sóng (tần số, hệ số trải phổ, khoảng cách, góc phương vị, chênh "
    "lệch độ cao), nhóm đặc trưng địa hình tính từ bản đồ độ cao số (độ dốc, độ gồ ghề, "
    "thống kê độ cao dọc đường truyền) và nhóm đặc trưng vật cản (tỷ lệ vật cản trên "
    "elip Fresnel, khoảng cách thông thoáng nhỏ nhất và trung bình, tỷ lệ vùng dân cư). "
    "Bảng 5.1 liệt kê chi tiết các đặc trưng đầu vào.",
)

add_para(
    doc,
    "Bảng 5.1. Bảng đặc trưng đầu vào của mô hình học máy",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=11.5,
)

input_rows = [
    ("1", "Tần số sóng mang", "Hz", "Tần số kênh truyền (~923 MHz cho AS923-2)"),
    ("2", "Hệ số trải phổ", "Số nguyên", "Giá trị từ 7 đến 12"),
    ("3", "Log khoảng cách 2D", "log10(m)", "Logarit cơ số 10 của khoảng cách ngang"),
    ("4", "Log khoảng cách 3D", "log10(m)", "Đã kể đến chênh lệch độ cao"),
    ("5", "Chênh lệch vĩ độ", "Radian", "Toạ độ tương đối thiết bị – gateway"),
    ("6", "Chênh lệch kinh độ", "Radian", "Toạ độ tương đối thiết bị – gateway"),
    ("7", "Góc phương vị", "Radian", "Hướng từ gateway đến thiết bị"),
    ("8", "Độ cao gateway", "Mét", "Lấy từ bản đồ độ cao số"),
    ("9", "Chênh lệch độ cao", "Mét", "Bao gồm chiều cao anten"),
    ("10", "Góc ngẩng đường truyền", "Radian", "Góc nâng nhìn từ gateway đến thiết bị"),
    ("11", "Độ dốc địa hình", "Mét", "Phương sai độ cao cục bộ"),
    ("12", "Độ gồ ghề địa hình", "Mét", "Phương sai cửa sổ 5×5"),
    ("13", "Trung bình độ cao đường truyền", "Mét", "Tính dọc tuyến gateway – thiết bị"),
    ("14", "Phương sai độ cao đường truyền", "Mét", "Đo mức biến thiên địa hình"),
    ("15", "Độ cao nhỏ nhất trên đường truyền", "Mét", "Điểm thấp nhất"),
    ("16", "Độ cao lớn nhất trên đường truyền", "Mét", "Điểm cao nhất"),
    ("17", "Tỷ lệ vật cản Fresnel", "Tỷ lệ 0–1", "Số điểm bị chặn / tổng số điểm trên elip"),
    ("18", "Khoảng cách thông thoáng nhỏ nhất", "Mét", "Khe Fresnel hẹp nhất"),
    ("19", "Khoảng cách thông thoáng trung bình", "Mét", "Trung bình dọc đường truyền"),
    ("20", "Tỷ lệ vùng dân cư", "Tỷ lệ 0–1", "Tính từ dữ liệu sử dụng đất OpenStreetMap"),
    ("21", "Mã định danh gateway", "Phân loại", "Mã hoá một-nóng (one-hot)"),
]
add_table(
    doc,
    ["STT", "Tên đặc trưng", "Đơn vị", "Mô tả ngắn gọn"],
    input_rows,
    col_widths_cm=[1.2, 5.5, 2.8, 6.0],
)

add_para(doc, "")

# --- 1.2 Cach tinh dac trung dia hinh va vat can ---
add_heading(doc, "1.2. Cách tính các đặc trưng địa hình và vật cản", level=2)

add_para(
    doc,
    "Hai nhóm đặc trưng quan trọng nhất ảnh hưởng trực tiếp đến chất lượng dự đoán là "
    "nhóm địa hình và nhóm vật cản đường truyền. Phần này mô tả cách tính từng nhóm để "
    "bảo đảm khả năng tái lập của quy trình.",
)

add_para(
    doc,
    "Đặc trưng địa hình được tính từ bản đồ độ cao số (DEM) ghép từ hai nguồn: bộ dữ "
    "liệu Copernicus DEM độ phân giải 30 mét cho địa hình và bộ dữ liệu DSM ghép từ "
    "ESA WorldCover bổ sung vùng phủ thực vật, công trình. Hai bộ dữ liệu được nạp vào "
    "bộ nhớ qua thư viện rasterio và được truy vấn bởi tiện ích crc_covlib của tổ "
    "chức tiêu chuẩn ITU-R. Đối với mỗi cặp gateway – thiết bị, hệ thống lấy mẫu một "
    "trăm điểm cách đều dọc đường thẳng nối hai điểm và tra cứu giá trị độ cao tại "
    "từng điểm. Từ chuỗi độ cao này, các giá trị thống kê được tính gồm: trung bình, "
    "phương sai, cực tiểu, cực đại; độ dốc tổng được tính bằng độ lệch chuẩn của các "
    "khoảng độ cao kế tiếp; độ gồ ghề được tính trên cửa sổ 5×5 quanh mỗi điểm.",
)

add_para(
    doc,
    "Đặc trưng vật cản Fresnel mô tả mức độ thông thoáng của vùng không gian hình elip "
    "quanh đường truyền — yếu tố quyết định trong lý thuyết sóng vô tuyến. Bán kính "
    "elip Fresnel bậc nhất tại điểm cách gateway một khoảng d₁ và cách thiết bị một "
    "khoảng d₂ được tính theo công thức:",
)
add_para(doc, "    r₁ = √( λ · d₁ · d₂ / (d₁ + d₂) )", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
add_para(
    doc,
    "trong đó λ là bước sóng (xấp xỉ 0,325 mét ở tần số 923 MHz). Tại mỗi điểm trên "
    "đường truyền, hệ thống so sánh độ cao địa hình thực tế với mức an toàn (đường "
    "thẳng giữa hai anten trừ đi sáu mươi phần trăm bán kính Fresnel). Tỷ lệ điểm bị "
    "chặn so với tổng số điểm cho ra đặc trưng tỷ lệ vật cản, các khoảng cách thông "
    "thoáng nhỏ nhất và trung bình cũng được trích từ chuỗi sai khác này.",
)

add_para(
    doc,
    "Đặc trưng tỷ lệ vùng dân cư được tính từ dữ liệu sử dụng đất OpenStreetMap đã "
    "được nướng vào tệp ảnh phân loại. Tại mỗi vị trí thiết bị, hệ thống lấy hình "
    "tròn bán kính năm mươi mét, đếm số ô vuông được phân loại là khu dân cư và chia "
    "cho tổng số ô — kết quả nằm trong khoảng từ 0 đến 1.",
)

# --- 1.3 Phan tang du lieu chong ro ri ---
add_heading(doc, "1.3. Phân tầng dữ liệu chống rò rỉ", level=2)

add_para(
    doc,
    "Một mô hình học máy chỉ có giá trị thực tế nếu chỉ tiêu đánh giá đo trên dữ liệu "
    "mà nó chưa từng nhìn thấy. Trong các phiên bản đầu của hệ thống, dữ liệu huấn "
    "luyện và dữ liệu kiểm thử có thể nằm trong cùng một khu vực địa lý và cùng một "
    "phiên đo, dẫn đến hiện tượng rò rỉ dữ liệu — mô hình ghi nhớ vị trí cụ thể thay "
    "vì học quy luật suy hao truyền sóng, làm cho RMSE báo cáo thấp hơn thực tế. Để "
    "khắc phục, hệ thống áp dụng phân tầng theo ô lưới lục giác H3.",
)

add_para(
    doc,
    "Lưới H3 là một hệ thống lưới phân cấp toàn cầu do tổ chức Uber phát triển: bề "
    "mặt Trái Đất được chia thành các ô lục giác đều nhau, mỗi ô được gán một mã "
    "định danh duy nhất. Mỗi mức phân giải H3 ứng với một kích thước ô khác nhau. "
    "Bảng 5.2 trình bày so sánh ba mức phân giải được xem xét.",
)

add_para(
    doc,
    "Bảng 5.2. So sánh ba mức phân giải H3",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=11.5,
)
add_table(
    doc,
    ["Mức", "Diện tích mỗi ô", "Quy mô tương đương", "Đánh giá"],
    [
        (
            "7",
            "≈ 5,16 km²",
            "Một quận nhỏ",
            "Quá thô — ô kiểm thử trùm nhiều khu khác nhau, mất chi tiết",
        ),
        ("8", "≈ 0,74 km²", "Một khu phố", "Phù hợp mật độ dữ liệu Đà Nẵng — đã chọn"),
        (
            "9",
            "≈ 0,10 km²",
            "Một khối nhà",
            "Quá mịn — nhiều ô chỉ có một mẫu, không đủ phương sai",
        ),
    ],
    col_widths_cm=[1.2, 3.0, 3.6, 8.4],
)
add_para(doc, "")

add_para(
    doc,
    "Trong mỗi ô H3, các bản ghi của cùng một thiết bị trong cùng một khung giờ được "
    "gom thành một phiên đo. Cửa sổ một giờ được chọn vì với tốc độ đi bộ trung bình "
    "năm kilômét trên giờ, người khảo sát có thể di chuyển một khoảng cách nhỏ hơn "
    "đường kính của một ô H3 mức 8 trong một giờ — bảo đảm các bản ghi cùng phiên "
    "đều nằm trong cùng một ô lưới.",
)

add_para(
    doc,
    "Quy tắc phân tầng được áp dụng như sau: các phiên mới nhất được chọn vào tập "
    "kiểm thử cho đến khi đủ định mức một nghìn năm trăm bản ghi, kế đến các phiên "
    "tiếp theo được chọn vào tập xác thực với định mức tương tự, phần còn lại thuộc "
    "tập huấn luyện. Sau bước này, hệ thống kiểm tra hai điều kiện bắt buộc bằng "
    "câu lệnh khẳng định:",
)
add_para(
    doc,
    "    Tập hợp ô H3 của tập huấn luyện và tập hợp ô H3 của tập kiểm thử không có "
    "phần tử chung (cell-disjoint). Tập hợp định danh phiên đo của tập huấn luyện "
    "và tập kiểm thử cũng không có phần tử chung (session-disjoint).",
    italic=True,
)
add_para(
    doc,
    "Nếu một trong hai điều kiện không thoả, quá trình huấn luyện dừng lại và báo lỗi "
    "ngay tại bước chuẩn bị dữ liệu — bảo đảm các con số báo cáo không bị nhiễm rò "
    "rỉ. Cơ chế phân tầng được minh hoạ trong Hình 5.3.",
)

add_image(
    doc, "hinh_5_3.png", "Hình 5.3. Phân tầng ô lưới H3 cho ba tập huấn luyện, xác thực và kiểm thử"
)

add_para(
    doc,
    "Hình 5.2 trực quan hoá phân bố dữ liệu trên hai trục đặc trưng quan trọng nhất "
    "là khoảng cách và tỷ lệ vật cản, kèm theo giá trị RSSI tương ứng. Có thể thấy "
    "phần lớn dữ liệu tập trung trong khoảng 0 – 6 km và RSSI giảm rõ rệt khi khoảng "
    "cách hoặc tỷ lệ vật cản tăng.",
)

add_image(doc, "hinh_5_2.png", "Hình 5.2. Phân bố dữ liệu trên hai trục đặc trưng quan trọng")

# --- 1.4 Kien truc hoc phan du hai tang ---
add_heading(doc, "1.4. Kiến trúc học phần dư hai tầng", level=2)

add_para(
    doc,
    "Hệ thống sử dụng kiến trúc hai tầng kết hợp kiến thức vật lý truyền sóng với "
    "khả năng học từ dữ liệu thực. Tầng thứ nhất là mô hình vật lý theo khuyến nghị "
    "ITU-R P.1812 (mô hình suy hao đường truyền cho dải tần 30 MHz – 6 GHz) bổ sung "
    "thêm thành phần suy hao vật cản theo khuyến nghị ITU-R P.2108-1. Tầng này sử "
    "dụng bản đồ độ cao số làm đầu vào và cho ra giá trị RSSI cơ sở dựa hoàn toàn "
    "trên định luật vật lý.",
)

add_para(
    doc,
    "Tầng thứ hai là mô hình học máy Extra Trees. Khác với cách tiếp cận end-to-end "
    "(học máy dự đoán trực tiếp RSSI), tầng này được huấn luyện để dự đoán phần dư "
    "Δ giữa giá trị RSSI đo được trong thực tế và giá trị RSSI mà tầng thứ nhất tính "
    "ra. Giá trị RSSI cuối cùng phục vụ ứng dụng là tổng của hai tầng:",
)
add_para(
    doc, "    RSSI cuối = RSSI Tầng 1 + Δ", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, bold=True
)

add_para(
    doc,
    "Lợi ích của kiến trúc hai tầng gồm: tận dụng được hiểu biết vật lý đã được "
    "kiểm chứng (mô hình ITU-R đã được tiêu chuẩn hoá quốc tế), giảm gánh nặng cho "
    "tầng học máy (chỉ cần học phần sai lệch nhỏ so với cơ sở vật lý chứ không phải "
    "toàn bộ quy luật suy hao), và bảo đảm dự đoán vẫn hợp lý ngay cả ở các vùng "
    "xa nguồn dữ liệu khảo sát — nơi tầng học máy có thể chưa đủ thông tin nhưng "
    "tầng vật lý vẫn hoạt động.",
)

add_para(
    doc,
    "Cơ chế dự phòng được tích hợp ở tầng thứ hai: trong trường hợp tệp mô hình "
    "joblib bị hỏng hoặc dịch vụ ml-service không sẵn sàng, hệ thống đặt Δ = 0 và "
    "trả về kết quả của tầng thứ nhất — bảo đảm hệ thống không bị gián đoạn dịch "
    "vụ trong mọi trường hợp. Sơ đồ tổng thể của kiến trúc hai tầng được trình "
    "bày trong Hình 5.5.",
)

add_image(
    doc,
    "hinh_5_5.png",
    "Hình 5.5. Kiến trúc học phần dư hai tầng — kết hợp ITU-R P.1812 và Extra Trees",
)

# --- 1.5 Huan luyen Extra Trees ---
add_heading(doc, "1.5. Huấn luyện thuật toán Extra Trees", level=2)

add_para(
    doc,
    "Thuật toán Extra Trees Regressor được lựa chọn nhờ ba ưu điểm chính: tốc độ huấn "
    "luyện nhanh nhờ chia ngưỡng ngẫu nhiên, khả năng chống quá khớp tốt hơn rừng ngẫu "
    "nhiên truyền thống và không cần tinh chỉnh siêu tham số phức tạp. Cấu hình siêu "
    "tham số sử dụng trong lần huấn luyện chính được liệt kê trong Bảng 5.3.",
)

add_para(
    doc,
    "Bảng 5.3. Cấu hình siêu tham số của Extra Trees",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=11.5,
)
add_table(
    doc,
    ["Siêu tham số", "Giá trị", "Ý nghĩa"],
    [
        ("Số cây", "1 500", "Số bộ ước lượng cơ sở"),
        ("Độ sâu tối đa", "20", "Giới hạn chiều sâu mỗi cây để chống quá khớp"),
        ("Số mẫu tối thiểu để tách nút", "5", "Ngưỡng tách nhánh"),
        ("Số mẫu tối thiểu tại nút lá", "2", "Số mẫu tối thiểu của một nút lá"),
        ("Hạt giống ngẫu nhiên", "42", "Bảo đảm khả năng tái lập"),
        ("Song song hoá", "Tất cả lõi", "Tận dụng toàn bộ CPU khi huấn luyện"),
    ],
    col_widths_cm=[5.0, 3.0, 7.5],
)
add_para(doc, "")

add_para(
    doc,
    "Đường ống huấn luyện gồm ba bước: chuẩn hoá đặc trưng số (điền giá trị khuyết bằng "
    "trung vị rồi đưa về phân phối chuẩn), mã hoá đặc trưng phân loại (mã hoá một-nóng "
    "đối với định danh gateway) và huấn luyện Extra Trees trên ma trận đặc trưng kết "
    "hợp. Sau khi huấn luyện xong, mô hình được lưu thành tệp joblib và được sao chép "
    "nguyên tử ra vị trí dùng cho dịch vụ phục vụ.",
)

# --- 1.6 Quy trinh huan luyen lai tu dong ---
add_heading(doc, "1.6. Quy trình huấn luyện lại tự động", level=2)

add_para(
    doc,
    "Để duy trì chất lượng mô hình theo thời gian khi dữ liệu khảo sát mới được cộng "
    "đồng đóng góp, hệ thống cung cấp quy trình huấn luyện lại tự động được điều phối "
    'qua hàng đợi Celery. Khi quản trị viên bấm nút "Huấn luyện lại" trên giao diện '
    "điều hành, dịch vụ api-service phát một công việc vào hàng đợi và ghi sự kiện "
    "vào bảng audit.ml_retrain_jobs. Tiến trình Celery worker lấy công việc và chạy "
    "tuần tự bốn bước con như một chuỗi tiến trình con độc lập. Bảng 5.4 trình bày "
    "thời gian giới hạn cho từng bước.",
)

add_para(
    doc,
    "Bảng 5.4. Bốn bước trong quy trình huấn luyện lại tự động",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=11.5,
)
add_table(
    doc,
    ["Bước", "Tệp kịch bản", "Thời gian giới hạn", "Đầu ra chính"],
    [
        (
            "1. Tạo CSV huấn luyện",
            "build_training_csv.py",
            "40 phút",
            "Tệp devices_history_full.csv kèm cột data_split",
        ),
        (
            "2. Huấn luyện Extra Trees",
            "train_extra_trees.py",
            "60 phút",
            "Tệp mô hình joblib (đổi tên nguyên tử .new → .joblib)",
        ),
        (
            "3. Đánh giá tập kiểm thử",
            "eval_extra_trees_holdout.py",
            "5 phút",
            "Tệp holdout_eval.json + per_distance + per_gateway",
        ),
        (
            "4. Xuất báo cáo",
            "render_ml_report.py",
            "10 phút",
            "Tệp report.pdf kèm các biểu đồ và bảng tổng hợp",
        ),
    ],
    col_widths_cm=[3.8, 4.0, 2.5, 5.5],
)
add_para(doc, "")

add_para(
    doc,
    "Ba điểm đáng chú ý trong cách thiết kế quy trình. Thứ nhất, sau bước thứ hai, "
    "tệp mô hình mới được ghi với phần mở rộng tạm thời .new rồi mới đổi tên thành "
    "tệp chính thức — kỹ thuật đổi tên nguyên tử của hệ điều hành bảo đảm dịch vụ "
    "phục vụ không bao giờ đọc phải tệp đang ghi dở. Thứ hai, sau khi đổi tên xong, "
    "dịch vụ ml-service được kích hoạt lại nóng (hot-reload) qua lời gọi POST "
    "/admin/reload kèm mã thông báo Bearer — bộ trọng số mới được nạp vào bộ nhớ mà "
    "không cần khởi động lại tiến trình, không gây gián đoạn dịch vụ. Thứ ba, kết "
    "quả của từng bước (số mẫu, các chỉ tiêu RMSE/MAE/R²/Bias, định danh người gọi) "
    "được ghi vào audit log để phục vụ kiểm tra về sau.",
)

add_para(
    doc,
    "Cơ chế chịu lỗi mềm được áp dụng ở bước kích hoạt lại nóng: nếu lời gọi tới "
    "ml-service thất bại, hệ thống ghi cảnh báo và để mô hình cũ tiếp tục phục vụ "
    "thay vì làm hỏng hệ thống. Quản trị viên có thể can thiệp thủ công khi cần. "
    "Sơ đồ chuỗi bốn bước được trình bày trong Hình 5.6.",
)

add_image(doc, "hinh_5_6.png", "Hình 5.6. Sơ đồ chuỗi bốn bước huấn luyện lại tự động qua Celery")

# --- 1.7 Trien khai ---
add_heading(doc, "1.7. Triển khai và phục vụ dự đoán", level=2)

add_para(
    doc,
    "Mô hình đã huấn luyện được triển khai trong dịch vụ ml-service (FastAPI). Khi có "
    "yêu cầu dự đoán, dịch vụ ml-service trích đặc trưng đầu vào từ toạ độ và cấu hình "
    "thiết bị, gọi mô hình Extra Trees để dự đoán phần dư so với mô hình vật lý ITU-R "
    "P.1812 (Tầng 1) và trả về phần dư đó cho dịch vụ api-service tổng hợp. Sau mỗi lần "
    "huấn luyện lại, dịch vụ ml-service được kích hoạt lại nóng (hot-reload) bằng lời "
    "gọi điểm cuối /admin/reload — tệp joblib mới được nạp lại trong bộ nhớ mà không "
    "cần khởi động lại dịch vụ. Bảng 5.4 mô tả các trường đầu ra của một lần gọi dự đoán.",
)

add_para(
    doc,
    "Bảng 5.5. Bảng đầu ra của một lần dự đoán",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=11.5,
)
add_table(
    doc,
    ["STT", "Trường đầu ra", "Đơn vị", "Mô tả ngắn gọn"],
    [
        ("1", "RSSI dự đoán", "dBm", "Cường độ tín hiệu nhận được sau hiệu chỉnh"),
        ("2", "SNR dự đoán", "dB", "Tỷ số tín hiệu trên nhiễu"),
        (
            "3",
            "Trạng thái phủ sóng",
            "Phân loại",
            "Một trong bốn mức: rất tốt, tốt, kém, không phủ",
        ),
        ("4", "Gateway phục vụ", "Mã định danh", "Gateway có liên kết tốt nhất"),
        ("5", "Độ tin cậy dự đoán", "Tỷ lệ 0–1", "Mức tin cậy ước lượng"),
        ("6", "Phiên bản mô hình", "Chuỗi", "Định danh phiên bản mô hình đang phục vụ"),
        ("7", "Đường lên (uplink)", "Đối tượng", "RSSI, SNR và biên độ dự phòng theo hướng lên"),
        (
            "8",
            "Đường xuống (downlink)",
            "Đối tượng",
            "RSSI, SNR và biên độ dự phòng theo hướng xuống",
        ),
    ],
    col_widths_cm=[1.2, 4.5, 2.8, 7.0],
)

add_para(doc, "")

# ============================================================================
# Muc 2 - Danh gia mo hinh hoc may
# ============================================================================
add_heading(doc, "2. Đánh giá mô hình học máy", level=1)

add_para(
    doc,
    "Mô hình được đánh giá độc lập trên ba tập: tập huấn luyện (đánh giá nội tại), "
    "tập xác thực (theo dõi quá khớp) và tập kiểm thử (đánh giá khả năng tổng quát "
    "hoá ngoài phân phối). Bốn chỉ tiêu đánh giá được sử dụng là: sai số bình phương "
    "trung bình căn bậc hai (RMSE), sai số tuyệt đối trung bình (MAE), hệ số xác định "
    "(R²) và độ lệch hệ thống (bias). Kết quả tổng thể được tổng hợp trong Bảng 5.6.",
)

add_para(
    doc,
    "Bảng 5.6. Kết quả đánh giá mô hình trên ba tập dữ liệu",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=11.5,
)
add_table(
    doc,
    ["Tập", "Số mẫu", "RMSE (dB)", "MAE (dB)", "R²", "Bias (dB)"],
    [
        ("Huấn luyện", "14 017", "4,07", "2,46", "0,869", "−0,48"),
        ("Xác thực", "1 514", "7,38", "5,50", "0,197", "—"),
        ("Kiểm thử", "1 636", "6,32", "4,75", "−0,076", "−2,89"),
    ],
    col_widths_cm=[2.6, 2.2, 2.4, 2.4, 2.4, 2.4],
)

add_para(doc, "")
add_para(
    doc,
    "Để phân tích sâu hơn, kết quả trên tập kiểm thử được chia theo khoảng cách giữa "
    "thiết bị và gateway. Khoảng cách ngắn (dưới 5 km) cho sai số cao hơn so với khoảng "
    "cách xa do hiện tượng đa đường trong môi trường đô thị. Bảng 5.7 trình bày chi tiết.",
)

add_para(
    doc,
    "Bảng 5.7. RMSE trên tập kiểm thử chia theo khoảng cách",
    italic=True,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=11.5,
)
add_table(
    doc,
    ["Khoảng cách", "Số mẫu", "RMSE (dB)", "MAE (dB)", "Bias (dB)"],
    [
        ("0 – 2 km", "862", "7,11", "5,39", "−3,91"),
        ("2 – 5 km", "150", "9,26", "8,22", "−7,13"),
        ("5 – 10 km", "276", "3,77", "2,94", "+0,14"),
        ("10 – 50 km", "348", "3,83", "3,08", "−0,95"),
    ],
    col_widths_cm=[3.0, 2.4, 3.0, 3.0, 3.0],
)

add_para(doc, "")
add_image(doc, "hinh_5_4a.png", "Hình 5.4a. So sánh RMSE và MAE trên ba tập dữ liệu", width_cm=12.0)
add_image(
    doc, "hinh_5_4b.png", "Hình 5.4b. RMSE trên tập kiểm thử chia theo khoảng cách", width_cm=12.0
)

add_heading(doc, "2.1. Nhận xét kết quả", level=2)

add_para(
    doc,
    "RMSE trên tập kiểm thử đạt 6,32 dB. Con số này nằm trong vùng giới hạn vật lý của "
    "hiện tượng shadow fading ở môi trường đô thị (theo khuyến nghị ITU-R P.1812 dao "
    "động trong khoảng 6 – 8 dB) và thấp hơn 4,26 dB so với mô hình XGBoost phiên bản "
    "trước (10,58 dB). So với mô hình truyền thống COST-231 / Hata thường cho sai số "
    "trên 20 dB ở khu vực đô thị, mức cải thiện là rất rõ rệt.",
)
add_para(
    doc,
    "Khoảng cách trong nhóm 2 – 5 km cho sai số cao nhất (9,26 dB) cùng độ lệch hệ "
    "thống −7,13 dB. Đây là vùng chuyển tiếp giữa đô thị lõi và vùng ven, có nhiều "
    "vật cản nhỏ rời rạc khó đặc trưng hoá. Ngược lại, ở khoảng cách trên 5 km, "
    "đường truyền chủ yếu là tầm nhìn thẳng hoặc gần thẳng nên sai số chỉ còn 3,77 – "
    "3,83 dB.",
)
add_para(
    doc,
    "Hệ số xác định R² trên tập kiểm thử mang giá trị âm (−0,076). Đây không phải dấu "
    "hiệu mô hình kém: do tập kiểm thử được phân tầng cell-disjoint, phương sai RSSI "
    "trong bảy ô kiểm thử nhỏ hơn nhiều so với phương sai toàn tập, dẫn đến R² âm khi "
    "chuẩn hoá theo phương sai cục bộ. Trong trường hợp này, RMSE và MAE phản ánh "
    "trung thực hơn chất lượng dự đoán.",
)
add_para(
    doc,
    "Độ lệch hệ thống −2,89 dB cho thấy mô hình có xu hướng dự đoán cao hơn so với "
    "giá trị đo được. Bảy ô kiểm thử phần nhiều rơi vào khu vực vật cản dày hơn mặt "
    "bằng chung của tập huấn luyện. Đây là đặc thù khi phân tầng cell-disjoint với số "
    "lượng ô kiểm thử nhỏ và có thể được giảm bớt bằng cách bổ sung khảo sát ở các "
    "khu vực còn thiếu dữ liệu.",
)

add_heading(doc, "2.2. Hướng cải tiến", level=2)

add_para(
    doc,
    "Trên cơ sở các phân tích trên, hai hướng cải tiến chính được đề xuất. Thứ nhất, "
    "bổ sung dữ liệu khảo sát ở các vùng còn ít mẫu, đặc biệt là vùng chuyển tiếp đô "
    "thị – ngoại thành (2 – 5 km) và các vùng ven biển. Thứ hai, làm phong phú thêm "
    "đặc trưng đầu vào bằng cách bổ sung chiều cao thực tế của nhà cửa từ dữ liệu "
    "vector, mật độ thảm thực vật ở mức chi tiết hơn, và phân biệt môi trường trong "
    "nhà / ngoài trời. Ngoài ra, có thể thực hiện điều chỉnh siêu tham số bằng tìm "
    "kiếm có hệ thống để giảm chênh lệch giữa tập huấn luyện và tập kiểm thử.",
)

# ============================================================================
# Luu file
# ============================================================================
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"Saved: {OUT}")
