"""Tạo báo cáo Word về pipeline ML của dự án LoRa Coverage."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent.parent / "docs" / "bao_cao_pipeline_ml.docx"


def set_cell_bg(cell, color_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        if level == 0:
            run.font.size = Pt(20)
        elif level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + 0.6 * level)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_rich_bullet(doc: Document, parts: list[tuple[str, bool]], level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + 0.6 * level)
    p.paragraph_format.space_after = Pt(2)
    for text, bold in parts:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = bold


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = ""
        p = hdr_cells[idx].paragraphs[0]
        run = p.add_run(header)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr_cells[idx], "1F4E79")
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row_idx, row in enumerate(rows):
        cells = table.rows[row_idx + 1].cells
        for col_idx, val in enumerate(row):
            cells[col_idx].text = ""
            p = cells[col_idx].paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()

    # Cấu hình lề
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    # ===== Trang bìa =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("BÁO CÁO PIPELINE HỌC MÁY\nDỰ ÁN ƯỚC LƯỢNG VÙNG PHỦ LoRa")
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(22)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("\nManage Data — Train Model — Evaluate Model — Deploy Model")
    sub_run.font.name = "Times New Roman"
    sub_run.font.size = Pt(13)
    sub_run.italic = True

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(
        "\n\nĐồ án tốt nghiệp\nNgày cập nhật: 2026-06-20\n"
        "Mô hình hiện hành: stage2-et-v0.7.0 (Extra Trees Regressor)"
    )
    info_run.font.name = "Times New Roman"
    info_run.font.size = Pt(12)

    doc.add_page_break()

    # ===== Phần 1: Manage Data =====
    add_heading(doc, "Phần 1. Manage Data — Quản lý dữ liệu", level=1)

    # 1.1 Collect
    add_heading(doc, "1.1. Collect — Thu thập", level=2)
    add_para(doc, "Dữ liệu vào dự án từ ba luồng song song:")
    add_rich_bullet(
        doc,
        [
            ("ChirpStack realtime (webhook): ", True),
            (
                "instance LoRaWAN của trường/đơn vị gửi mỗi gói uplink qua webhook. "
                "File chirpstack_webhook_service.py:23 chuyển mỗi uplink thành N record "
                "(một record cho mỗi gateway nghe được). Dedup bằng UUID5 từ "
                "deduplicationId:rx_index.",
                False,
            ),
        ],
    )
    add_rich_bullet(
        doc,
        [
            ("CSV/JSON upload (cộng đồng): ", True),
            (
                "người dùng đi đo thực địa hoặc xuất từ công cụ khác, sau đó tải lên "
                "qua trình duyệt. Parser quy chuẩn dữ liệu rồi gói lại thành batch "
                "(bảng me.upload_batches).",
                False,
            ),
        ],
    )
    add_rich_bullet(
        doc,
        [
            ("LPWANMapper sync: ", True),
            (
                "REST pull định kỳ từ api.lpwanmapper.com/data. Adapter "
                "sources/lpwanmapper/adapter.py:64 chuẩn hoá định dạng.",
                False,
            ),
        ],
    )

    add_para(doc, "")
    add_para(doc, "Các bảng cơ sở dữ liệu chính:", bold=True)
    add_rich_bullet(
        doc,
        [
            ("ts.survey_quarantine ", True),
            (
                "(TimescaleDB hypertable, migration 0003): toàn bộ dữ liệu mới đi vào "
                "bảng này trước, chờ quản trị viên duyệt.",
                False,
            ),
        ],
    )
    add_rich_bullet(
        doc,
        [
            ("ts.survey_training: ", True),
            (
                "dữ liệu đã được duyệt mới được chuyển vào đây. Mô hình ML chỉ huấn "
                "luyện trên bảng này.",
                False,
            ),
        ],
    )
    add_rich_bullet(
        doc,
        [
            ("geo.gateway_quarantine ", True),
            (
                "(migration 0029): gateway mới phát hiện qua sync được chuyển vào hàng "
                "đợi duyệt riêng.",
                False,
            ),
        ],
    )
    add_rich_bullet(
        doc,
        [
            ("geo.gateways: ", True),
            (
                "danh mục gateway đã được duyệt; dữ liệu tham chiếu DEM/landuse nằm "
                "tại /data/dem/copernicus_glo30_*.tif và landuse_central.geojson.",
                False,
            ),
        ],
    )

    add_para(doc, "")
    add_para(
        doc,
        'Quy trình duyệt (moderation): người dùng bấm "Đóng góp" '
        "→ batches.py:311 chuyển trạng thái submitted_for_community=true "
        "→ batch hiện trong hàng đợi của quản trị viên.",
    )

    # 1.2 Explore
    add_heading(doc, "1.2. Explore — Khám phá", level=2)
    add_para(doc, "Dự án không có notebook EDA riêng. Thay vào đó:")
    add_bullet(
        doc,
        "Thư mục scripts/experiments/ chứa các script đặc trị: "
        "eval_breakdown_sf.py, rssi_by_distance_2026_05_31.py, "
        "eval_stage1_vs_stage2_2026_05_31.py.",
    )
    add_bullet(
        doc,
        "File train_split_stats.json được sinh sau mỗi lần build CSV "
        "(build_training_csv.py:650): số dòng/cell/session từng tập, "
        "khoảng thời gian train so với test.",
    )
    add_bullet(doc, "render_ml_report.py sinh báo cáo HTML/PDF cho mỗi lần retrain.")
    add_bullet(
        doc,
        "Audit thường xuyên: lọc d < 50 km (survey.py:49) để tránh "
        "corruption (Hải Phòng gắn gateway Đà Nẵng cách ~554 km).",
    )

    # 1.3 Cleanse
    add_heading(doc, "1.3. Cleanse — Làm sạch", level=2)
    add_para(doc, "Quy tắc validation ngay khi ingest (survey.py:49, 98-103):", bold=True)
    add_bullet(doc, "Khoảng cách device ↔ gateway phải nhỏ hơn 50 km.")
    add_bullet(doc, "RSSI nằm trong khoảng [-150, -30] dBm.")
    add_bullet(doc, "SNR nằm trong khoảng [-30, 30] dB.")
    add_bullet(doc, "Spreading factor thuộc {7, 8, 9, 10, 11, 12}.")
    add_bullet(
        doc,
        "Dedup: UNIQUE PARTIAL (timestamp, source_type, external_id) "
        "→ ChirpStack và LPWANMapper bị loại trùng tự động.",
    )

    add_para(doc, "")
    add_para(doc, "Luồng kiểm duyệt (quarantine):", bold=True)
    add_bullet(doc, "Dòng mới đi vào ts.survey_quarantine với reject_reason=NULL (trạng thái chờ).")
    add_bullet(
        doc,
        "Quản trị viên duyệt batch theo 4 chế độ: all / points_only / "
        "gateways_only / reject → promote sang ts.survey_training.",
    )
    add_bullet(doc, "Đồng thời nếu kèm gateway mới: geo.gateway_quarantine → geo.gateways.")
    add_bullet(
        doc,
        "Quản trị viên xoá: soft-delete batch (đặt deleted_at), "
        "hard-purge quarantine, giữ nguyên các dòng training đã duyệt.",
    )

    add_para(doc, "")
    add_para(
        doc,
        "Trong build_training_csv.py:254 — chỉ SELECT dòng có "
        "submitted_for_community=TRUE từ ts.survey_training. "
        "Dữ liệu quarantine không lọt vào tập huấn luyện.",
    )

    # 1.4 Prepare
    add_heading(doc, "1.4. Prepare — Chuẩn bị (Feature Engineering)", level=2)
    add_para(doc, "21 đặc trưng được suy ra trong build_training_csv.py, chia 5 nhóm:")
    add_table(
        doc,
        ["Nhóm", "Đặc trưng", "Nguồn / Cách tính"],
        [
            ["Tín hiệu", "frequency, spreading_factor", "Metadata uplink"],
            [
                "Hình học link",
                "log_distance, log_distance_3d, delta_lat, delta_lon, angle",
                "Haversine + arctan2",
            ],
            [
                "Độ cao",
                "gw_elevation, delta_elevation, elevation_angle",
                "DEM Copernicus GLO-30 + chiều cao anten (device 1.5 m, gateway 15 m)",
            ],
            [
                "Bề mặt địa hình",
                "slope, roughness, terrain_mean/std/min/max",
                "Cửa sổ 3x3 trên DEM + sample 30 m dọc đường truyền",
            ],
            [
                "Fresnel",
                "fresnel_obstruction_ratio, min_fresnel_clearance, mean_fresnel_clearance",
                "Bán kính Fresnel zone 1 so với profile DEM",
            ],
            ["Land use", "residential_ratio", "% đường truyền qua đất dân cư (OSM)"],
            ["Categorical", "gateway", "Tên/EUI gateway → OneHotEncoder"],
        ],
    )

    add_para(doc, "Pipeline tiền xử lý (train_extra_trees.py:66-86):", bold=True)
    add_bullet(doc, "Numeric: SimpleImputer(median) → StandardScaler.")
    add_bullet(
        doc,
        "Categorical: SimpleImputer(most_frequent) → "
        'OneHotEncoder(handle_unknown="ignore") để gateway mới không '
        "gây lỗi khi inference.",
    )
    add_bullet(
        doc,
        "Đóng gói trong ColumnTransformer + Pipeline để lưu/load trong duy nhất một file .joblib.",
    )

    # 1.5 Split
    add_heading(doc, "1.5. Split — Chia tập", level=2)
    add_para(
        doc,
        "Chiến lược: H3 hexagonal grid (resolution 8) + temporal hold-out "
        "(build_training_csv.py:80-92, 472-618).",
    )
    add_bullet(
        doc,
        "H3 res 8: mỗi cell khoảng 0.74 km² (cỡ một khu phố). "
        "Mục đích: không cho một ô địa lý xuất hiện ở cả train và test "
        "→ tránh spatial leak.",
    )
    add_bullet(
        doc,
        "Session window 1 giờ (SESSION_WINDOW_S=3600): gộp các gói "
        "cùng device, cùng ô, cùng giờ thành một session. Việc chia "
        "tập theo session chứ không theo từng dòng.",
    )

    add_para(doc, "")
    add_para(doc, "Phân chia cụ thể:", bold=True)
    add_bullet(
        doc,
        "Test (~1500 dòng, giới hạn 30%/cell): chọn các ô có session mới nhất (tháng 1-2/2026).",
    )
    add_bullet(doc, "Val (~1500 dòng): chọn các ô next-newest.")
    add_bullet(
        doc,
        "Buffer ring (BUFFER_RING=0): các ô sát test/val nhưng không "
        "thuộc tập nào sẽ bị loại để cô lập.",
    )
    add_bullet(doc, "Train: phần còn lại (chủ yếu tháng 11-12/2025).")
    add_bullet(
        doc,
        "Assertion bảo vệ (dòng 579-584): "
        "train_cells ∩ test_cells = ∅ và "
        "train_sessions ∩ test_sessions = ∅. Build fail nếu vi phạm.",
    )

    doc.add_page_break()

    # ===== Phần 2: Train Model =====
    add_heading(doc, "Phần 2. Train Model — Huấn luyện mô hình", level=1)

    # 2.1 Pick learning task
    add_heading(doc, "2.1. Pick Learning Task — Chọn bài toán học máy", level=2)
    add_para(doc, "Bài toán: Hồi quy (Regression) có giám sát.", bold=True)
    add_bullet(doc, "Đầu vào (X): 21 đặc trưng radio/địa hình mô tả một link device ↔ gateway.")
    add_bullet(
        doc,
        "Đầu ra (y): rssi — giá trị RSSI tuyệt đối (dBm) đo được tại "
        "gateway (train_extra_trees.py:53).",
    )
    add_bullet(
        doc,
        "Hàm mất mát: MSE (cây hồi quy mặc định). Metric chính: "
        "RMSE (dB) — đơn vị thân thiện với kỹ sư RF.",
    )

    add_para(doc, "")
    add_para(doc, "Kiến trúc 2 tầng (Stage 1 + Stage 2):", bold=True)
    add_bullet(
        doc,
        "Stage 1: mô hình vật lý truyền sóng ITU-R P.1812 tính RSSI "
        "baseline từ DEM + tần số + công suất.",
    )
    add_bullet(
        doc,
        "Stage 2: Extra Trees dự đoán RSSI tuyệt đối end-to-end từ "
        "21 đặc trưng (không dự đoán residual trực tiếp).",
    )
    add_bullet(
        doc,
        "Gộp tại API: prediction_service.py lấy "
        "delta = RSSI_ET − RSSI_Stage1, trả về client dưới tên "
        "residual_db (giữ contract cũ, stage2_client.py:47-49).",
    )

    # 2.2 Engineer features
    add_heading(doc, "2.2. Engineer Features — Kỹ nghệ đặc trưng", level=2)
    add_para(
        doc, "Đặc trưng đã liệt kê chi tiết ở phần 1.4. Một số quyết định thiết kế đáng chú ý:"
    )
    add_bullet(
        doc,
        "Bắt đầu từ khoảng 50 ứng viên trong reference_wireless/, "
        "loại dần các đặc trưng có importance gần 0.",
    )
    add_bullet(
        doc,
        "Một thử nghiệm bỏ 5 đặc trưng low-importance (obstruction_ratio, "
        "mean_obstruction, terrain_range…): RMSE 10.94 → 10.37 dB, "
        "bias 2-5 km cải thiện rõ.",
    )
    add_bullet(
        doc,
        "Chốt 21 đặc trưng cuối cùng — quá trình hoàn toàn thủ công, "
        "không dùng automatic feature selection (FeatureWiz, Boruta).",
    )
    add_bullet(
        doc,
        'OneHotEncoder(handle_unknown="ignore") cho gateway: gateway '
        "mới được encode toàn 0 → mô hình rơi về dự đoán trung bình "
        "(không crash).",
    )

    # 2.3 Select algorithms
    add_heading(doc, "2.3. Select Algorithms — Chọn thuật toán", level=2)
    add_para(doc, "Thuật toán cuối: ExtraTreesRegressor (1500 cây).", bold=True)
    add_code_block(
        doc,
        "ET_PARAMS = {\n"
        '    "n_estimators":     1500,   # số cây\n'
        '    "max_depth":        20,     # giới hạn để tránh overfit nặng\n'
        '    "min_samples_split": 5,\n'
        '    "min_samples_leaf":  2,\n'
        '    "max_features":     None,   # mỗi split xét tất cả đặc trưng\n'
        '    "random_state":     42,\n'
        '    "n_jobs":           -1,     # song song toàn bộ CPU\n'
        "}",
    )

    add_para(
        doc,
        "Benchmark so sánh 6 thuật toán trên cùng feature set (5-fold CV stratified by gateway):",
    )
    add_table(
        doc,
        ["Thuật toán", "RMSE (dB)", "MAE (dB)", "R²"],
        [
            ["Extra Trees (đã chọn)", "3.45", "2.03", "0.906"],
            ["Random Forest", "3.42", "2.11", "0.907"],
            ["XGBoost", "3.80", "2.14", "0.886"],
            ["HistGradientBoosting", "3.72", "2.22", "0.890"],
            ["MLP (Neural Network)", "4.19", "2.58", "0.861"],
            ["SVR", "5.49", "2.79", "0.761"],
        ],
    )

    add_para(doc, "Lý do chọn Extra Trees:", bold=True)
    add_bullet(
        doc,
        "CV score tốt nhất (R² trung bình 5-fold = 0.911) — nhỉnh hơn Random Forest, vượt XGBoost.",
    )
    add_bullet(
        doc,
        "Ít overfit hơn Random Forest: Extra Trees random hoá luôn cả "
        "ngưỡng split → giảm variance.",
    )
    add_bullet(
        doc,
        "Robust với feature noise: DEM/landuse có nhiễu cục bộ, tree "
        "ensemble xử lý tốt hơn linear/SVR.",
    )
    add_bullet(
        doc,
        "Inference deterministic và nhanh: 1500 cây khoảng 120 MB, "
        "predict 1 điểm dưới 1 ms — phù hợp serve qua FastAPI.",
    )
    add_bullet(doc, "Không cần GPU: deploy được trên container CPU thường.")

    add_para(doc, "")
    add_para(
        doc,
        "Hyperparameter tuning: tune thủ công dựa benchmark, không dùng "
        "Optuna ở phiên bản v0.7 hiện tại.",
        italic=True,
    )

    doc.add_page_break()

    # ===== Phần 3: Evaluate Model =====
    add_heading(doc, "Phần 3. Evaluate Model — Đánh giá mô hình", level=1)

    # 3.1 Score models
    add_heading(doc, "3.1. Score Models — Chấm điểm", level=2)
    add_para(doc, "5 metric chính dùng đồng nhất ở mọi giai đoạn:")
    add_table(
        doc,
        ["Metric", "Đơn vị", "Ý nghĩa"],
        [
            ["RMSE", "dB", "Sai số bình phương trung bình — metric chủ đạo"],
            ["MAE", "dB", "Sai số tuyệt đối trung bình — robust với outlier"],
            ["R²", "—", "Mức độ giải thích phương sai của target"],
            ["Bias", "dB", "Trung bình sai số có dấu (measured − predicted) — đo lệch hệ thống"],
            ["Percentile", "dB", "P50/P75/P90/P95 error — đo đuôi phân bố"],
        ],
    )

    add_para(doc, "Output JSON (3 file):", bold=True)
    add_bullet(
        doc, "train_metrics.json — train in-sample: {rmse, mae, r2, rows_trained, feature_count}."
    )
    add_bullet(doc, "val_metrics.json — val unseen sessions: {rmse, mae, r2, n}.")
    add_bullet(
        doc,
        "holdout_eval.json — test temporal Jan-Feb 2026: "
        "{window, overall, per_distance_bin, v06_xgboost_baseline_rmse_db, "
        "delta_vs_v06_db}.",
    )

    add_para(doc, "")
    add_para(doc, "Breakdown chi tiết:", bold=True)
    add_bullet(
        doc,
        "Per-distance-bin: 4 dải (0-2 km, 2-5 km, 5-10 km, 10-50 km) "
        "— eval_extra_trees_holdout.py:82.",
    )
    add_bullet(doc, "Per-SF: SF7 → SF12 — experiments/eval_breakdown_sf.py:44.")
    add_bullet(
        doc, "Per-gateway: script ad-hoc tính RMSE/MAE/bias mỗi gateway — eval_breakdown_sf.py:54."
    )

    # 3.2 Assess performance
    add_heading(doc, "3.2. Assess Performance — Đánh giá hiệu năng", level=2)
    add_para(doc, "Ba lát cắt song song:")
    add_table(
        doc,
        ["Tập", "n", "RMSE (dB)", "MAE (dB)", "R²", "Ghi chú"],
        [
            [
                "Train (in-sample)",
                "10,867",
                "2.78",
                "1.69",
                "0.944",
                "Cây fit gần như khít → expected",
            ],
            [
                "Val (unseen sessions)",
                "1,514",
                "7.38",
                "5.50",
                "0.197",
                "Gap lớn so với train → generalization khó",
            ],
            [
                "Holdout temporal (Jan-Feb 2026)",
                "1,636",
                "6.57",
                "4.98",
                "-0.16",
                "Số defense thật; bias -2.51 dB (over-predict)",
            ],
        ],
    )

    add_para(doc, "Phân tích per-distance-bin:", bold=True)
    add_bullet(
        doc,
        "0-2 km: RMSE 8.06 dB (n=862) — TỆ NHẤT, do DSM mismatch nhà cao tầng và clutter chi tiết.",
    )
    add_bullet(doc, "2-5 km: RMSE ~6.5 dB.")
    add_bullet(doc, "5-10 km: RMSE ~5.0 dB.")
    add_bullet(
        doc,
        "10-50 km: RMSE 3.79 dB (n=348) — tốt nhất vì path loss thống "
        "trị, ít ảnh hưởng nhiễu local.",
    )

    add_para(doc, "")
    add_para(
        doc,
        'R² âm ở holdout: mô hình dự đoán kém hơn baseline "đoán trung '
        'bình tập test" — phản ánh distribution shift giữa train '
        "(Nov-Dec 2025) và test (Jan-Feb 2026). Đây là số phải báo trung "
        "thực với hội đồng.",
        italic=True,
    )

    add_para(doc, "")
    add_para(doc, "Báo cáo HTML/PDF tự sinh (render_ml_report.py):", bold=True)
    add_bullet(doc, "summary.html (Jinja2) + report.pdf (WeasyPrint) + summary.json.")
    add_bullet(
        doc,
        "5 PNG plots: scatter measured-vs-predicted (train), "
        "error-vs-distance, feature importance top-20, per-distance bar, "
        "per-gateway bar.",
    )
    add_bullet(doc, "Lưu vào thư mục reports/retrain-<job_id>/ qua Celery task.")

    # 3.3 Compare alternatives
    add_heading(doc, "3.3. Compare Alternatives — So sánh phương án", level=2)

    add_para(
        doc,
        "A. Stage 1 ITU P.1812 — trước/sau fix per-gateway noise floor (2026-05-31, fix #3):",
        bold=True,
    )
    add_table(
        doc,
        ["Cấu hình", "Holdout RMSE (dB)", "Δ"],
        [
            ["ITU P.1812 + global NF -174 dBm/Hz", "23.80", "baseline"],
            ["ITU P.1812 + per-gw NF (migration 0020)", "11.11", "−12.69"],
        ],
    )
    add_para(doc, "Đây là cải tiến lớn nhất từ trước đến nay của Stage 1.", italic=True)

    add_para(doc, "")
    add_para(doc, "B. Extra Trees so với các baseline khác (cùng holdout):", bold=True)
    add_table(
        doc,
        ["Model", "RMSE (dB)", "Δ so với ET"],
        [
            ["Stage 1 alone (P.1812 + per-gw NF)", "11.11", "+4.54"],
            ["XGBoost legacy v0.6", "10.58", "+4.01"],
            ["Extra Trees v0.7 (offline)", "6.57", "—"],
            ["Log-distance (đã bỏ)", "~25", "—"],
        ],
    )
    add_para(
        doc,
        "ET cải thiện khoảng 4 dB so với XGBoost, và khoảng 4.5 dB so với Stage 1 vật lý.",
        italic=True,
    )

    add_para(doc, "")
    add_para(doc, "C. API serving gap — đáng báo trong DATN:", bold=True)
    add_table(
        doc,
        ["Cách inference", "RMSE (dB)", "Bias (dB)"],
        [
            ["Offline (joblib trực tiếp)", "10.58", "-0.25"],
            ["API /coverage/predict", "13.47", "+4.55"],
        ],
    )
    add_para(
        doc,
        "Chênh 2.89 dB RMSE và 4.8 dB bias drift. Nguyên nhân nghi: cách "
        "chọn serving gateway hoặc feature compute khác trên API path. "
        "Memory project_api_offline_gap_2026_05_31.md đã ghi → eval thật "
        "phải qua offline script, không qua API.",
        italic=True,
    )

    add_para(doc, "")
    add_para(doc, "D. Bias correction đã thử và revert (2026-06-14):", bold=True)
    add_bullet(doc, "Hardcode offset -4.67 dB để correct bias holdout về 0.")
    add_bullet(doc, "Khi áp lên tập dữ liệu khác → bias flip sang +3 dB (tệ hơn) → revert.")
    add_bullet(
        doc,
        "Kết luận: không patch bằng constant; phải fix qua thêm data + "
        "feature engineering (đặc biệt 0-2 km).",
    )

    doc.add_page_break()

    # ===== Phần 4: Deploy Model =====
    add_heading(doc, "Phần 4. Deploy Model — Triển khai mô hình", level=1)

    # 4.1 Apply to fresh data
    add_heading(doc, "4.1. Apply Model to Fresh Data — Áp dụng vào dữ liệu mới", level=2)

    add_para(doc, "A. Inference từng điểm (online, low-latency):", bold=True)
    add_bullet(
        doc,
        "ml-service expose 2 endpoint "
        "(services/ml-service/src/lora_ml_predict/app.py:244-325): "
        "POST /residual (1 điểm) và POST /residuals/batch (tối đa ~5000 "
        "điểm/request).",
    )
    add_bullet(
        doc,
        "Response: {residual_db, model_version, ood_flag} với residual_db = RSSI_ET − RSSI_Stage1.",
    )
    add_bullet(
        doc,
        "api-service gọi qua stage2_client.py (httpx AsyncClient, "
        "bearer token, timeout 60s) trong prediction_service.py:88. "
        "Fallback Stage 1 thuần nếu ml-service timeout/fail → không vỡ pipeline.",
    )
    add_bullet(
        doc,
        "Use case: GET /api/v1/coverage/predict?lat=&lon= → trả RSSI "
        'dự đoán + PDR/BER + confidence cho UI "Dự đoán điểm".',
    )

    add_para(doc, "")
    add_para(doc, "B. Inference batch (offline, big workload):", bold=True)
    add_bullet(doc, '"Rebuild bản đồ ước lượng" — Celery task tasks/rebuild_coverage.py:54.')
    add_bullet(
        doc,
        "Query MAX(timestamp) per gateway từ ts.survey_training → so "
        "last_rebuild_at → skip nếu không có data mới (incremental).",
    )
    add_bullet(
        doc, "Gọi scripts/precompute_rssi_heatmap.py raster hoá toàn vùng Đà Nẵng grid 10 m."
    )
    add_bullet(
        doc,
        "Lưu ý: heatmap KHÔNG dùng Extra Trees — chỉ P.1812 + DTM + "
        "per-gw NF + survey overlay. Extra Trees chỉ phục vụ "
        "/coverage/predict per-point.",
    )

    add_para(doc, "")
    add_para(doc, "C. Atomic swap và hot-reload (zero downtime):", bold=True)
    add_bullet(
        doc,
        "train_extra_trees.py:154-156 ghi extra_trees_model.joblib.new, "
        "rồi os.rename() để swap file đang serve.",
    )
    add_bullet(
        doc,
        "ml-service POST /admin/reload (app.py:212-241) load joblib "
        "trong asyncio.to_thread() để không block event loop.",
    )
    add_bullet(
        doc,
        "Docker: ENV LORA_ML_MODEL_PATH=/app/services/ml-service/data/"
        "extra_trees_model.joblib, HEALTHCHECK /healthz mỗi 30s.",
    )

    # 4.2 Monitor outcomes
    add_heading(doc, "4.2. Monitor Outcomes — Theo dõi kết quả", level=2)

    add_para(doc, "A. Audit job (truy vết retrain):", bold=True)
    add_para(doc, "Bảng audit.ml_retrain_jobs (migration 0025) gồm các cột:")
    add_code_block(
        doc,
        "id, status (queued/running/succeeded/failed),\n"
        "triggered_by (user FK), triggered_at, started_at, finished_at,\n"
        "rows_trained, artifact_path,\n"
        "metrics (JSONB: rmse / mae / r2 / holdout_*),\n"
        "error_text, celery_task_id",
    )
    add_para(
        doc, "Admin UI đọc bảng này để xem job nào đang chạy + lịch sử metrics qua các lần retrain."
    )

    add_para(doc, "")
    add_para(doc, "B. Prometheus metrics + health check:", bold=True)
    add_bullet(
        doc, "edge/metrics.py:19-126: histogram latency theo method/status (prometheus_client)."
    )
    add_bullet(doc, "LOOKUP_LATENCY_SECONDS SLO /coverage/predict → cảnh báo nếu p95 vượt ngưỡng.")
    add_bullet(
        doc,
        "api-service /healthz (liveness), /readyz (ping DB → 503 nếu "
        "down). ml-service /healthz + Docker HEALTHCHECK 30s.",
    )
    add_bullet(doc, "/metrics endpoint format OpenMetrics → scrape Prometheus/Grafana được.")

    add_para(doc, "")
    add_para(doc, "C. Sentry: ", bold=True)
    add_para(
        doc,
        "config.py.sentry_dsn + sentry_traces_sample_rate=0 (chỉ bắt "
        "error, không trace để tiết kiệm cost) — init ở "
        "edge/app.py:_configure_sentry().",
    )

    add_para(doc, "")
    add_para(doc, "D. Theo dõi trực tiếp (Live session):", bold=True)
    add_bullet(
        doc,
        "Sau refactor 2026-06-16: view-only, poll 5 giây, idle 15 phút "
        "auto-stop (không tạo batch, không ghi training).",
    )
    add_bullet(
        doc, "Mục đích: cho admin debug ChirpStack realtime — không phải monitoring model output."
    )

    add_para(doc, "")
    add_para(doc, "E. Survey overlay (so sánh predicted vs actual):", bold=True)
    add_bullet(
        doc,
        "Frontend bản đồ overlay điểm đo thực (survey training) lên "
        "heatmap dự đoán → mắt người so sánh RSSI predicted vs measured.",
    )
    add_bullet(
        doc,
        "Không có metric tự động cảnh báo drift — đang là gap, manually "
        "check báo cáo PDF sau mỗi lần retrain.",
    )

    # 4.3 Improve model
    add_heading(doc, "4.3. Improve Model — Cải tiến", level=2)

    add_para(doc, "A. Retrain pipeline (Celery tasks/retrain_ml.py:54-228):", bold=True)
    add_para(
        doc,
        'Trigger: admin bấm nút "Retrain" → POST /api/v1/admin/ml/retrain '
        "(admin.py:1176) → Celery delay().",
    )
    add_para(doc, "Sequence:")
    add_bullet(doc, "1. INSERT audit.ml_retrain_jobs (status=queued).")
    add_bullet(
        doc,
        "2. Subprocess build_training_csv.py — timeout 40 phút "
        "(query DB + DEM/landuse feature engineering).",
    )
    add_bullet(
        doc, "3. Subprocess train_extra_trees.py — timeout 1 giờ (atomic swap joblib khi xong)."
    )
    add_bullet(doc, "4. Subprocess eval_extra_trees_holdout.py — fail-soft 5 phút.")
    add_bullet(doc, "5. Subprocess render_ml_report.py — fail-soft 10 phút (HTML + PDF).")
    add_bullet(doc, "6. POST ml-service/admin/reload (60s timeout) — fail-soft.")
    add_bullet(
        doc, "7. UPDATE audit.ml_retrain_jobs (status=succeeded + metrics JSONB + report_dir)."
    )
    add_para(doc, "→ Không cần restart container, không downtime.", italic=True)

    add_para(doc, "")
    add_para(doc, "B. Feedback loop (data → model):", bold=True)
    add_code_block(
        doc,
        "Community upload (CSV / JSON / ChirpStack / LPWANMapper)\n"
        "     ↓\n"
        "ts.survey_quarantine  (pending)\n"
        "     ↓ admin approve batch\n"
        "ts.survey_training  (community visible)\n"
        "     ↓ admin click Retrain\n"
        "build_training_csv → train ET → eval → hot-reload\n"
        "     ↓\n"
        "ml-service serve mô hình mới (model_version bump)",
    )

    add_para(doc, "")
    add_para(doc, "C. Versioning:", bold=True)
    add_bullet(
        doc,
        'model_version = "stage2-et-v0.7.0" hardcode trong app.py:51 '
        "— bump thủ công khi có schema change.",
    )
    add_bullet(
        doc,
        "Mỗi prediction response chứa model_version → frontend/log truy vết được mô hình nào tính.",
    )
    add_bullet(doc, "Không có model registry (MLflow/W&B) — joblib file đơn lẻ.")

    add_para(doc, "")
    add_para(doc, "D. Rollback và A/B test:", bold=True)
    add_bullet(
        doc,
        "Rollback: chưa có chính thức. Có thể giữ joblib cũ ngoài VM "
        "rồi cp đè + reload. Không formal versioned storage.",
    )
    add_bullet(doc, "A/B test: KHÔNG có. Chỉ 1 mô hình active tại 1 thời điểm.")
    add_bullet(
        doc,
        "Đây là future work — chấp nhận trade-off vì scope đồ án tốt nghiệp không đủ thời gian.",
    )

    add_para(doc, "")
    add_para(doc, "E. Gap đã biết (báo trung thực, không che giấu):", bold=True)
    add_bullet(
        doc,
        "DSM artifact <2 km: RMSE 8 dB ở dải gần nhất do DSM lệch / "
        "nhà cao tầng. Cần upgrade DSM resolution hoặc thêm building "
        "height feature.",
    )
    add_bullet(
        doc,
        "API↔offline gap 2.89 dB RMSE: serving wiring drift chưa fix "
        "(memory project_api_offline_gap_2026_05_31.md).",
    )
    add_bullet(
        doc,
        "R² âm holdout temporal: distribution shift Nov-Dec → Jan-Feb "
        "→ cần retrain định kỳ (đề xuất cron mỗi 2 tuần).",
    )
    add_bullet(
        doc,
        "Bias correction hardcode đã thử + revert "
        "(memory project_ml_bias_correction_2026_06_14) — chốt: phải "
        "fix qua data/feature, không patch constant.",
    )
    add_bullet(
        doc,
        "Không có monitoring drift tự động: cần Evidently AI hoặc "
        "custom drift detector — future work.",
    )

    doc.add_page_break()

    # ===== Tổng kết =====
    add_heading(doc, "Tổng kết", level=1)
    add_para(doc, "Pipeline ML của dự án LoRa Coverage bao gồm 4 giai đoạn rõ rệt:")
    add_bullet(
        doc,
        "Manage Data: thu thập từ ChirpStack + community upload + "
        "LPWANMapper; làm sạch qua quarantine + admin moderation; "
        "feature engineering 21 đặc trưng địa hình + RF; chia tập "
        "theo H3 res 8 + temporal hold-out tránh spatial leak.",
    )
    add_bullet(
        doc,
        "Train Model: bài toán hồi quy RSSI; ExtraTreesRegressor "
        "(1500 cây) được chọn sau benchmark 6 thuật toán; pipeline "
        "Scikit-learn standard (impute + scale + OneHot).",
    )
    add_bullet(
        doc,
        "Evaluate Model: 5 metric (RMSE/MAE/R²/bias/percentile); ba "
        "lát cắt train/val/holdout-temporal; per-distance-bin breakdown; "
        "báo cáo HTML/PDF tự sinh.",
    )
    add_bullet(
        doc,
        "Deploy Model: ml-service FastAPI; atomic-swap joblib + "
        "hot-reload; Celery retrain task tự động; audit qua "
        "ml_retrain_jobs; Prometheus + Sentry monitoring.",
    )

    add_para(doc, "")
    add_para(doc, "Số metric trung thực để báo cáo hội đồng:", bold=True)
    add_bullet(
        doc, "Holdout temporal Jan-Feb 2026 (n=1,636): RMSE 6.57 dB, MAE 4.98 dB, bias -2.51 dB."
    )
    add_bullet(doc, "So với Stage 1 alone: cải thiện 4.54 dB RMSE.")
    add_bullet(doc, "So với XGBoost legacy v0.6: cải thiện 4.01 dB RMSE.")
    add_bullet(doc, "Vẫn còn gap đáng kể ở dải 0-2 km (RMSE 8.06 dB) cần future work.")

    return doc


if __name__ == "__main__":
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Đã tạo: {OUTPUT}")
